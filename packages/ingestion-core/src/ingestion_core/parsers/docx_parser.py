from __future__ import annotations

from docx import Document

from ingestion_core.parsers.base import BaseParser, ParseOutput

MAX_TEXT_CHARS = 120_000


class DocxParser(BaseParser):
    def parse(self, path: str) -> ParseOutput:
        doc = Document(path)
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for t_i, table in enumerate(doc.tables):
            parts.append(f"[table {t_i + 1}]")
            for row in table.rows:
                parts.append(" | ".join(c.text.strip() for c in row.cells))
        text = "\n".join(parts)[:MAX_TEXT_CHARS]
        return ParseOutput(
            text=text, structured_hint="Word document.",
            metadata={"tables": len(doc.tables)},
        )
